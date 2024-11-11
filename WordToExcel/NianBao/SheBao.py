import os
import xlwt
from WordToExcel.path_tool import get_word_folder_path, get_output_folder_path
from docx import Document


# 创建excel文档以及表头
workbook = xlwt.Workbook(encoding='utf-8')
sheet = workbook.add_sheet('年报-社保信息')
title = ['序号', '企业序号', '统一社会信用代码', '年份',
         '城镇职工基本养老保险', '职工基本医疗保险', '生育保险', '失业保险', '工伤保险',
         '单位参加城镇职工基本养老保险缴费基数', '单位参加失业保险缴费基数', '单位参加职工基本医疗保险缴费基数', '单位参加生育保险缴费基数',
         '单位参加城镇职工基·······································································································································本养老保险本期实际缴费金额', '单位参加失业保险本期实际缴费金额', '单位参加职工基本医疗保险本期实际缴费金额', '单位参加职工基本医疗保险本期实际缴费金额', '单位参加生育保险本期实际缴费金额',
         '单位参加城镇职工基本养老保险累计欠费金额', '单位参加失业保险累计欠费金额', '单位参加职工基本医疗保险累计欠费金额', '单位参加工伤保险累计欠费金额', '单位参加生育保险累计欠费金额'
         ]
# 序号
row_num = 0
for col, column in enumerate(title):
    sheet.write(row_num, col, column)
row_num += 1

# 文件夹路径，包含所有.docx文件
word_folder_path = get_word_folder_path()
# 文件夹路径，包含所有.cls文件
output_folder = get_output_folder_path()



# 企业序号
qiye_id = 0
# 遍历文件夹中的所有.docx文件
for filename in os.listdir(word_folder_path):
    if filename.endswith('.docx'):


        # word的绝对路径
        docx_path = os.path.join(word_folder_path, filename)

        # 创建Document对象
        doc = Document(docx_path)

        # 统一社会信用代码(第一个表格的 第四行 第四列)
        if doc.tables[0].rows[3].cells[2].text == "统一社会信用代码":
            credit_code = doc.tables[0].rows[3].cells[3].text
        else:
            credit_code = None

        # 写入excel
        tables = [table for table in doc.tables if table.rows[0].cells[0].text == "城镇职工基本养老保险"]

        if len(tables) != 0:
            qiye_id += 1

            for table_id, table in enumerate(tables):

                # 第一列，序号
                sheet.write(row_num, 0, row_num)
                # 第二列，企业序号（测试版本，到时候需要根据word来，一个word对应一个*********************）
                sheet.write(row_num, 1, qiye_id)
                # 第三列，统一社会信用代码
                sheet.write(row_num, 2, credit_code)
                # 第四列，年份
                if table_id == 0:
                    sheet.write(row_num, 3, '2023')
                if table_id == 1:
                    sheet.write(row_num, 3, '2022')
                if table_id == 2:
                    sheet.write(row_num, 3, '2021')
                # 第五列，城镇职工基本养老保险
                sheet.write(row_num, 4, table.rows[0].cells[1].text)
                # 第六列，职工基本医疗保险
                sheet.write(row_num, 5, table.rows[0].cells[3].text)
                # 第七列，生育保险
                sheet.write(row_num, 6, table.rows[1].cells[1].text)
                # 第八列，失业保险
                sheet.write(row_num, 7, table.rows[1].cells[3].text)
                # 第九列，工伤保险
                sheet.write(row_num, 8, table.rows[2].cells[1].text)

                # 第十列，单位参加城镇职工基本养老保险缴费基数
                sheet.write(row_num, 9, table.rows[3].cells[3].text)
                # 第十一列，单位参加失业保险缴费基数
                sheet.write(row_num, 10, table.rows[4].cells[3].text)
                # 第十二列，单位参加职工基本医疗保险缴费基数
                sheet.write(row_num, 11, table.rows[5].cells[3].text)
                # 第十三列，单位参加生育保险缴费基数
                sheet.write(row_num, 12, table.rows[6].cells[3].text)

                # 第十四列，单位参加城镇职工基本养老保险本期实际缴费金额
                sheet.write(row_num, 13, table.rows[7].cells[3].text)
                # 第十十五列，单位参加失业保险本期实际缴费金额
                sheet.write(row_num, 14, table.rows[8].cells[3].text)
                # 第十六列，单位参加职工基本医疗保险本期实际缴费金额
                sheet.write(row_num, 15, table.rows[9].cells[3].text)
                # 第十七列，单位参加职工基本医疗保险本期实际缴费金额
                sheet.write(row_num, 16, table.rows[10].cells[3].text)
                # 第十八列，单位参加生育保险本期实际缴费金额
                sheet.write(row_num, 17, table.rows[11].cells[3].text)

                # 第十九列，单位参加城镇职工基本养老保险累计欠费金额
                sheet.write(row_num, 18, table.rows[12].cells[3].text)
                # 第二十列，单位参加失业保险累计欠费金额
                sheet.write(row_num, 19, table.rows[13].cells[3].text)
                # 第二十一列，单位参加职工基本医疗保险累计欠费金额
                sheet.write(row_num, 20, table.rows[14].cells[3].text)
                # 第二十二列，单位参加工伤保险累计欠费金额
                sheet.write(row_num, 21, table.rows[15].cells[3].text)
                # 第二十三列，单位参加生育保险累计欠费金额
                sheet.write(row_num, 22, table.rows[16].cells[3].text)


                # 行号+ 1
                row_num += 1





cls_path = os.path.join(output_folder, '年报-社保信息.xls')

workbook.save(cls_path)