import os
import xlwt
from WordToExcel.path_tool import get_word_folder_path, get_output_folder_path
from docx import Document


# 创建excel文档以及表头
workbook = xlwt.Workbook(encoding='utf-8')
sheet = workbook.add_sheet('年报-企业资产状况信息')
title = ['序号', '企业序号', '统一社会信用代码', '年份', '资产总额', '浙江省财政厅', '营业总收入', '利润总额', '营业总收入中主营业务收入', '净利润', '纳税总额', '负债总额']
# 序号
row_num = 0
for col, column in enumerate(title):
    sheet.write(row_num, col, column)
row_num += 1

# 文件夹路径，包含所有.docx文件
word_folder_path = get_word_folder_path()
# 文件夹路径，包含所有.cls文件
output_folder = get_output_folder_path()



# 企业序号
qiye_id = 0
# 遍历文件夹中的所有.docx文件
for filename in os.listdir(word_folder_path):
    if filename.endswith('.docx'):
        print(filename)
        qiye_id += 1
        # word的绝对路径
        docx_path = os.path.join(word_folder_path, filename)

        # 创建Document对象
        doc = Document(docx_path)

        # 统一社会信用代码(第一个表格的 第四行 第四列)
        if doc.tables[0].rows[3].cells[2].text == "统一社会信用代码":
            credit_code = doc.tables[0].rows[3].cells[3].text
        else:
            credit_code = None
        # 写入excel
        tables = [table for table in doc.tables if table.rows[0].cells[0].text == "资产总额" and table.rows[3].cells[2].text == "负债总额"]

        for table_id, table in enumerate(tables):

            # 第一列，序号
            sheet.write(row_num, 0, row_num)
            # 第二列，企业序号（测试版本，到时候需要根据word来，一个word对应一个*********************）
            sheet.write(row_num, 1, qiye_id)
            # 第三列，统一社会信用代码
            sheet.write(row_num, 2, credit_code)
            # 第四列，年份
            if table_id == 0:
                sheet.write(row_num, 3, '2023')
            if table_id == 1:
                sheet.write(row_num, 3, '2022')
            if table_id == 2:
                sheet.write(row_num, 3, '2021')
            # 第五列，资产总额
            sheet.write(row_num, 4, table.rows[0].cells[1].text)
            # 第六列，浙江省财政厅
            sheet.write(row_num, 5, table.rows[0].cells[3].text)
            # 第七列，营业总收入
            sheet.write(row_num, 6, table.rows[1].cells[1].text)
            # 第八列，利润总额
            sheet.write(row_num, 7, table.rows[1].cells[3].text)
            # 第九列，营业总收入中主营业务收入
            sheet.write(row_num, 8, table.rows[2].cells[1].text)
            # 第十列，净利润
            sheet.write(row_num, 9, table.rows[2].cells[3].text)
            # 第十一列，纳税总额
            sheet.write(row_num, 10, table.rows[3].cells[1].text)
            # 第十二列，负债总额
            sheet.write(row_num, 11, table.rows[3].cells[3].text)


            # 行号+ 1
            row_num += 1





cls_path = os.path.join(output_folder, '年报-企业资产状况信息.xls')

workbook.save(cls_path)