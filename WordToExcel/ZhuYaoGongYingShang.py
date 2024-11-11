import os
import xlwt
from path_tool import get_word_folder_path, get_output_folder_path
from docx import Document

# 创建excel文档以及表头
workbook = xlwt.Workbook(encoding='utf-8')
sheet = workbook.add_sheet('主要供应商')
title = ['序号', '企业序号', '统一社会信用代码', '供应商序号', '供应商名称', '关联关系', '采购额（元）', '占比（%）']
# 序号
row_num = 0
for col, column in enumerate(title):
    sheet.write(row_num, col, column)
row_num += 1

# 文件夹路径，包含所有.docx文件
word_folder_path = get_word_folder_path()
# 文件夹路径，包含所有.cls文件
output_folder = get_output_folder_path()


# 企业序号
qiye_id = 0
# 遍历文件夹中的所有.docx文件
for filename in os.listdir(word_folder_path):
    if filename.endswith('.docx'):
        qiye_id += 1

        # word的绝对路径
        docx_path = os.path.join(word_folder_path, filename)

        # 创建Document对象
        doc = Document(docx_path)

        # 统一社会信用代码(第一个表格的 第四行 第四列)
        if doc.tables[0].rows[3].cells[2].text == "统一社会信用代码":
            credit_code = doc.tables[0].rows[3].cells[3].text
        else:
            credit_code = None

        # 写入excel
        for table in doc.tables:
            # 找到对应的table
            if table.rows[0].cells[1].text == "供应商名称":

                for hang_id, row in enumerate(table.rows):

                    # 不要word中table的表头信息(前面已经手动生成)
                    if hang_id != 0:

                        # 第一列，序号
                        sheet.write(row_num, 0, row_num)
                        # 第二列，企业序号（测试版本，到时候需要根据word来，一个word对应一个*********************）
                        sheet.write(row_num, 1, qiye_id)
                        # 第三列，统一社会信用代码
                        sheet.write(row_num, 2, credit_code)

                        # 其余列，纯专利信息（插入时列号+3，因为excel相比于word中的table，前面加了3列）
                        for lie_id, cell in enumerate(row.cells):
                            sheet.write(row_num, lie_id + 3, cell.text)
                            # print(cell.text)
                        # 行号+ 1
                        row_num += 1

                # 找到对应table后就停止
                break


            # # 如果当前企业的word中没用该栏目的表格信息
            # else:
            #
            #



cls_path = os.path.join(output_folder, '主要供应商.xls')

workbook.save(cls_path)